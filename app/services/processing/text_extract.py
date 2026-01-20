from typing import Any
from io import BytesIO
from docx import Document
import subprocess
import tempfile
from pathlib import Path
from app.services.gdrive.downloader import download_file_bytes, export_google_doc_bytes
from app.services.processing.pdf_extract import extract_text_from_pdf_bytes

DOC_MIME = "application/msword"
def _convert_doc_to_docx_bytes(doc_bytes: bytes) -> bytes:
    with tempfile.TemporaryDirectory() as td:
        td_path = Path(td)
        in_path = td_path / "input.doc"
        in_path.write_bytes(doc_bytes)

        subprocess.run(
            ["soffice", "--headless", "--nologo", "--convert-to", "docx", "--outdir", td, str(in_path)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
        )

        docx_files = list(td_path.glob("*.docx"))
        if not docx_files:
            raise ValueError("DOC to DOCX conversion failed (no output produced)")

        return docx_files[0].read_bytes()
GOOGLE_DOC_MIME = "application/vnd.google-apps.document"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

def get_text_for_drive_file(service: Any, file_meta: dict) -> str:
    file_id = file_meta["id"]
    mime_type = file_meta.get("mimeType") or ""

    if mime_type == GOOGLE_DOC_MIME:
        raw = export_google_doc_bytes(service, file_id, mime_type="text/plain")
        return raw.decode("utf-8", errors="ignore")

    if mime_type.startswith("text/"):
        raw = download_file_bytes(service, file_id)
        return raw.decode("utf-8", errors="ignore")


    if mime_type == "application/pdf":
        raw = download_file_bytes(service, file_id)
        return extract_text_from_pdf_bytes(raw)

    if mime_type == DOCX_MIME:
        raw = download_file_bytes(service, file_id)
        doc = Document(BytesIO(raw))

        parts: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)

        return "\n".join(parts).strip()

    if mime_type == DOC_MIME:
        raw = download_file_bytes(service, file_id)
        docx_bytes = _convert_doc_to_docx_bytes(raw)

        doc = Document(BytesIO(docx_bytes))
        parts: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)

        return "\n".join(parts).strip()

    raise ValueError(f"Unsupported mime type: {mime_type}")